# -*- coding: utf-8 -*-
"""轻量文档解析器 — 借鉴 StaffDeck knowledge/parser.py。

支持 txt/md/html/pdf/docx，有降级策略和编码检测。
不依赖 doc_processing 模块。
"""

from __future__ import annotations

from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}


class KnowledgeParseError(ValueError):
    """文档解析错误。"""


def extract_text(filename: str, content: bytes) -> tuple[str, str]:
    """提取文档文本。

    Args:
        filename: 文件名（用于判断格式）。
        content: 文件二进制内容。

    Returns:
        (text, format) 元组。

    Raises:
        KnowledgeParseError: 不支持的格式或解析失败。
    """
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise KnowledgeParseError(f"暂不支持 {suffix} 文件格式。")
    if suffix in {".txt", ".md", ".markdown"}:
        return _decode_text(content), suffix.lstrip(".")
    if suffix in {".html", ".htm"}:
        return _extract_html(content), "html"
    if suffix == ".pdf":
        return _extract_pdf(content), "pdf"
    if suffix == ".docx":
        return _extract_docx(content), "docx"
    raise KnowledgeParseError(f"暂不支持 {suffix} 文件格式。")


def _decode_text(content: bytes) -> str:
    """尝试多种编码解码文本。"""
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    """提取 PDF 文本。

    优先使用 pymupdf（项目已声明依赖），回退到 pypdf。
    """
    # 优先尝试 pymupdf（fitz）— 项目 pyproject.toml 已声明 pymupdf>=1.24.0
    try:
        import fitz  # pymupdf

        doc = fitz.open(stream=content, filetype="pdf")
        pages: list[str] = []
        for index, page in enumerate(doc):
            page_text = page.get_text() or ""
            if page_text.strip():
                pages.append(f"[Page {index + 1}]\n{page_text}")
        doc.close()
        if pages:
            return "\n\n".join(pages)
        # 文本层为空，返回提示
        return "(PDF 无可提取的文本层，可能是扫描件)"
    except Exception:
        pass

    # 回退：尝试 pypdf
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        pages = []
        for index, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(f"[Page {index + 1}]\n{page_text}")
        return "\n\n".join(pages) if pages else "(PDF 无可提取的文本层)"
    except Exception as exc:
        raise KnowledgeParseError(
            "PDF 解析失败：pymupdf 和 pypdf 均不可用。请安装 pymupdf (pip install pymupdf)。",
        ) from exc


def _extract_docx(content: bytes) -> str:
    """提取 DOCX 文本（优先 python-docx，降级 zip 解析）。"""
    try:
        from docx import Document

        document = Document(BytesIO(content))
        rows: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
        return "\n".join(rows)
    except Exception:
        return _extract_docx_with_zip(content)


def _extract_docx_with_zip(content: bytes) -> str:
    """降级：直接从 zip 读取 word/document.xml 提取文本。"""
    with ZipFile(BytesIO(content)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    parser = _DocxTextExtractor()
    parser.feed(xml)
    return parser.text


def _extract_html(content: bytes) -> str:
    """提取 HTML 文本（优先 BeautifulSoup，降级标准库）。"""
    text = _decode_text(content)
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(text, "html.parser")
        for item in soup(["script", "style", "noscript"]):
            item.decompose()
        return soup.get_text("\n")
    except Exception:
        parser = _HTMLTextExtractor()
        parser.feed(text)
        return parser.text


# ── 内部 HTML 解析器 ──────────────────────────────────────────────────────


class _HTMLTextExtractor(HTMLParser):
    """标准库 HTML 文本提取器（降级方案）。"""

    def __init__(self):
        super().__init__()
        self._parts: list[str] = []

    @property
    def text(self) -> str:
        return "\n".join(p.strip() for p in self._parts if p.strip())

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._parts.append(data)


class _DocxTextExtractor(HTMLParser):
    """从 DOCX XML 提取文本（降级方案）。"""

    def __init__(self):
        super().__init__()
        self._parts: list[str] = []

    @property
    def text(self) -> str:
        return "\n".join(p.strip() for p in self._parts if p.strip())

    def handle_data(self, data: str) -> None:
        if data.strip():
            self._parts.append(data)
